"""DOCX parser with image extraction and multi-modal support"""

import logging
import zipfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from xml.etree import ElementTree as ET

from docx import Document as DocxDocument
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from PIL import Image

from core.parsers.interfaces import (
    BaseParser,
    Document,
    DocumentMetadata,
    DocumentType,
    ParseError,
    Segment,
    SegmentType,
    TextSubtype,
    VisualSubtype,
    TableSubtype,
    VisualElement,
    VisualElementType,
)

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """
    DOCX parser with image extraction and multi-modal support
    
    Features:
    - Text extraction from Word documents
    - Image and diagram extraction
    - Table preservation
    - Multi-modal content with VLM descriptions
    - Paragraph-level segmentation
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None, enable_vlm: bool = True):
        """Initialize DOCX parser"""
        super().__init__(config, enable_vlm)
        self.supported_types = {DocumentType.DOCX}
        self.supported_extensions = {".docx", ".doc"}
        
        # DOCX-specific configuration
        self.config.setdefault("extract_images", True)
        self.config.setdefault("extract_tables", True)
        self.config.setdefault("preserve_formatting", False)
        self.config.setdefault("image_min_size", 50)
        self.config.setdefault("table_as_text", True)
        
        logger.info(f"Initialized DOCX parser with VLM: {enable_vlm}")
    
    def can_parse(self, file_path: Path) -> bool:
        """Check if file is a DOCX document"""
        return file_path.suffix.lower() in self.supported_extensions
    
    async def parse(self, file_path: Path) -> Document:
        """
        Parse DOCX document with image extraction
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Document with text segments and visual elements
            
        Raises:
            ParseError: If parsing fails
        """
        self.validate_file(file_path)
        
        try:
            logger.info(f"Starting DOCX parsing: {file_path.name}")
            
            # Load document
            doc = DocxDocument(str(file_path))
            
            # Extract metadata
            metadata = self._extract_docx_metadata(file_path, doc)
            
            # Extract visual elements
            visual_elements = []
            if self.config.get("extract_images", True):
                visual_elements = self._extract_visual_elements(file_path, doc)
            
            # Analyze visual elements with VLM
            if visual_elements and self.enable_vlm:
                document_context = {
                    "document_title": metadata.title,
                    "document_type": "DOCX",
                    "author": metadata.author
                }
                visual_elements = await self.analyze_visual_elements(visual_elements, document_context)
            
            # Create segments from document content
            segments = self._create_segments(doc, visual_elements)
            
            # Create visual segments
            visual_segments = self._create_visual_segments(visual_elements)
            
            # Merge all segments
            all_segments = segments + visual_segments
            
            # Re-index segments
            for i, segment in enumerate(all_segments):
                segment.segment_index = i
            
            # Create document
            document = self.create_document(
                file_path,
                segments=all_segments,
                metadata=metadata,
                visual_elements=visual_elements
            )
            
            logger.info(f"DOCX parsing completed: {len(segments)} segments, "
                       f"{len(visual_elements)} visual elements")
            
            return document
            
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}", exc_info=True)
            raise ParseError(f"Failed to parse DOCX {file_path.name}: {str(e)}")
    
    def _extract_docx_metadata(self, file_path: Path, doc: DocxDocument) -> DocumentMetadata:
        """Extract metadata from DOCX document"""
        try:
            # Get base metadata
            metadata = self.extract_metadata(file_path)
            metadata.document_type = DocumentType.DOCX
            
            # Extract document properties
            core_props = doc.core_properties
            if core_props:
                metadata.title = core_props.title or metadata.title
                metadata.author = core_props.author
                metadata.created_date = core_props.created
                metadata.modified_date = core_props.modified
                metadata.language = core_props.language
            
            # Count paragraphs as pages approximation
            metadata.page_count = len(doc.paragraphs)
            
            # Add custom metadata
            metadata.custom_metadata.update({
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections),
                "has_images": self._has_images(doc),
                "format": "DOCX"
            })
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {e}")
            metadata = self.extract_metadata(file_path)
            metadata.document_type = DocumentType.DOCX
            return metadata
    
    def _has_images(self, doc: DocxDocument) -> bool:
        """Check if document contains images"""
        try:
            for paragraph in doc.paragraphs:
                if self._paragraph_has_images(paragraph):
                    return True
            return False
        except Exception:
            return False
    
    def _paragraph_has_images(self, paragraph) -> bool:
        """Check if paragraph contains images"""
        try:
            for run in paragraph.runs:
                if run.element.xpath('.//pic:pic'):
                    return True
            return False
        except Exception:
            return False
    
    def _extract_visual_elements(self, file_path: Path, doc: DocxDocument) -> List[VisualElement]:
        """Extract visual elements from DOCX document"""
        visual_elements = []
        
        try:
            # Extract images using document structure
            visual_elements.extend(self._extract_embedded_images(file_path, doc))
            
            # Extract images from document relations
            visual_elements.extend(self._extract_relation_images(file_path))
            
            # Remove duplicates
            unique_elements = {}
            for element in visual_elements:
                if element.content_hash not in unique_elements:
                    unique_elements[element.content_hash] = element
            
            visual_elements = list(unique_elements.values())
            
            logger.info(f"Extracted {len(visual_elements)} visual elements from DOCX")
            return visual_elements
            
        except Exception as e:
            logger.error(f"Visual element extraction failed: {e}")
            return []
    
    def _extract_embedded_images(self, file_path: Path, doc: DocxDocument) -> List[VisualElement]:
        """Extract images embedded in document paragraphs"""
        visual_elements = []
        
        try:
            for para_idx, paragraph in enumerate(doc.paragraphs):
                for run in paragraph.runs:
                    # Find drawing elements in run
                    drawings = run.element.xpath('.//w:drawing')
                    
                    for drawing in drawings:
                        try:
                            # Extract image reference
                            blips = drawing.xpath('.//a:blip')
                            
                            for blip in blips:
                                embed_id = blip.get(qn('r:embed'))
                                if embed_id:
                                    # Get image data from document relationships
                                    image_data = self._get_image_data_from_embed(doc, embed_id)
                                    if image_data:
                                        # Determine image type
                                        img_type = self._determine_image_type(image_data)
                                        
                                        visual_element = VisualElement(
                                            element_type=img_type,
                                            source_format=DocumentType.DOCX,
                                            content_hash=VisualElement.create_hash(image_data),
                                            raw_data=image_data,
                                            page_or_slide=para_idx + 1,  # Use paragraph index as page
                                            segment_reference=f"paragraph_{para_idx}",
                                            file_extension=self._get_image_extension(image_data),
                                            analysis_metadata={
                                                "embed_id": embed_id,
                                                "paragraph_index": para_idx,
                                                "extraction_method": "embedded"
                                            }
                                        )
                                        
                                        visual_elements.append(visual_element)
                                        
                        except Exception as e:
                            logger.warning(f"Failed to extract embedded image: {e}")
                            continue
        
        except Exception as e:
            logger.error(f"Failed to extract embedded images: {e}")
        
        return visual_elements
    
    def _extract_relation_images(self, file_path: Path) -> List[VisualElement]:
        """Extract images from document relations"""
        visual_elements = []
        
        try:
            # Open DOCX as zip file
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # Find image files in media directory
                media_files = [f for f in zip_file.namelist() if f.startswith('word/media/')]
                
                for media_file in media_files:
                    try:
                        # Read image data
                        image_data = zip_file.read(media_file)
                        
                        # Skip very small images (likely icons)
                        if len(image_data) < 1000:
                            continue
                        
                        # Determine image type
                        img_type = self._determine_image_type(image_data)
                        
                        # Extract filename and extension
                        filename = Path(media_file).name
                        file_ext = Path(media_file).suffix.lower().lstrip('.')
                        
                        visual_element = VisualElement(
                            element_type=img_type,
                            source_format=DocumentType.DOCX,
                            content_hash=VisualElement.create_hash(image_data),
                            raw_data=image_data,
                            file_extension=file_ext,
                            analysis_metadata={
                                "media_file": media_file,
                                "filename": filename,
                                "extraction_method": "media_folder"
                            }
                        )
                        
                        visual_elements.append(visual_element)
                        
                    except Exception as e:
                        logger.warning(f"Failed to extract image from {media_file}: {e}")
                        continue
        
        except Exception as e:
            logger.error(f"Failed to extract relation images: {e}")
        
        return visual_elements
    
    def _get_image_data_from_embed(self, doc: DocxDocument, embed_id: str) -> Optional[bytes]:
        """Get image data from document relationships"""
        try:
            # Access document relationships
            part = doc.part
            if hasattr(part, 'rels') and embed_id in part.rels:
                rel = part.rels[embed_id]
                if hasattr(rel, 'target_part'):
                    return rel.target_part.blob
        except Exception as e:
            logger.debug(f"Failed to get image data for embed {embed_id}: {e}")
        
        return None
    
    def _determine_image_type(self, image_data: bytes) -> VisualElementType:
        """Determine the type of visual element"""
        try:
            # Check image content
            from PIL import Image
            import io
            
            img = Image.open(io.BytesIO(image_data))
            width, height = img.size
            
            # Heuristics for image type classification
            aspect_ratio = width / height
            
            # Very wide images might be charts
            if aspect_ratio > 2.5:
                return VisualElementType.CHART
            
            # Square-ish images might be diagrams
            if 0.8 <= aspect_ratio <= 1.2:
                return VisualElementType.DIAGRAM
            
            # Check for common chart/diagram characteristics
            # This is a simplified heuristic
            if width > 300 and height > 200:
                return VisualElementType.CHART
            
            return VisualElementType.IMAGE
            
        except Exception:
            return VisualElementType.UNKNOWN_VISUAL
    
    def _get_image_extension(self, image_data: bytes) -> str:
        """Get appropriate file extension for image"""
        try:
            # Check image format
            from PIL import Image
            import io
            
            img = Image.open(io.BytesIO(image_data))
            format_map = {
                'JPEG': 'jpg',
                'PNG': 'png',
                'BMP': 'bmp',
                'GIF': 'gif',
                'TIFF': 'tiff'
            }
            return format_map.get(img.format, 'png')
            
        except Exception:
            return 'png'  # Default
    
    def _create_segments(self, doc: DocxDocument, visual_elements: List[VisualElement]) -> List[Segment]:
        """Create text segments from document content"""
        segments = []
        
        try:
            # Process all paragraphs, including empty ones
            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text  # Don't strip to preserve empty paragraphs
                
                # Find visual references for this paragraph
                visual_refs = []
                for ve in visual_elements:
                    if ve.segment_reference == f"paragraph_{para_idx}":
                        visual_refs.append(ve.content_hash)
                
                # Determine segment type
                segment_type = self._determine_paragraph_type(paragraph)
                
                # For empty paragraphs, use a placeholder to satisfy validation
                # but mark them as empty in metadata
                content = text if text else " "  # Use single space for empty paragraphs
                
                # Create segment for all paragraphs
                segment = Segment(
                    content=content,
                    segment_type=SegmentType.TEXT,
                    segment_subtype=segment_type,  # segment_type from _determine_paragraph_type is actually subtype
                    segment_index=para_idx,  # Add proper segment index
                    visual_references=visual_refs,
                    metadata={
                        "paragraph_index": para_idx,
                        "style": paragraph.style.name if paragraph.style else None,
                        "alignment": str(paragraph.alignment) if paragraph.alignment else None,
                        "has_images": len(visual_refs) > 0,
                        "is_empty": len(text.strip()) == 0  # Mark empty paragraphs
                    }
                )
                segments.append(segment)
            
            # Process tables if enabled
            if self.config.get("extract_tables", True):
                segments.extend(self._process_tables(doc))
            
            logger.info(f"Created {len(segments)} segments from DOCX")
            return segments
            
        except Exception as e:
            logger.error(f"Segment creation failed: {e}")
            return []
    
    def _determine_paragraph_type(self, paragraph) -> str:
        """Determine the type of paragraph"""
        try:
            style_name = paragraph.style.name.lower() if paragraph.style else ""
            
            if "heading" in style_name:
                # Try to determine heading level
                if "1" in style_name:
                    return TextSubtype.HEADING_1.value
                elif "2" in style_name:
                    return TextSubtype.HEADING_2.value
                elif "3" in style_name:
                    return TextSubtype.HEADING_3.value
                else:
                    return TextSubtype.HEADING_1.value  # Default heading
            elif "title" in style_name:
                return TextSubtype.TITLE.value
            elif "subtitle" in style_name:
                return TextSubtype.SUBTITLE.value
            elif "caption" in style_name:
                return TextSubtype.CAPTION.value
            elif "quote" in style_name:
                return TextSubtype.QUOTE.value
            elif "list" in style_name:
                return TextSubtype.LIST.value
            else:
                return TextSubtype.PARAGRAPH.value
                
        except Exception:
            return TextSubtype.PARAGRAPH.value
    
    def _process_tables(self, doc: DocxDocument) -> List[Segment]:
        """Process tables in document"""
        table_segments = []
        
        try:
            for table_idx, table in enumerate(doc.tables):
                if self.config.get("table_as_text", True):
                    # Convert table to text representation
                    table_text = self._table_to_text(table)
                    if table_text:
                        segment = Segment(
                            content=table_text,
                            segment_type=SegmentType.TABLE,
                            segment_subtype=TableSubtype.DATA.value,
                            segment_index=len(table_segments),
                            metadata={
                                "table_index": table_idx,
                                "rows": len(table.rows),
                                "columns": len(table.columns) if table.rows else 0,
                                "extraction_method": "table_text"
                            }
                        )
                        table_segments.append(segment)
                        
        except Exception as e:
            logger.error(f"Table processing failed: {e}")
        
        return table_segments
    
    def _table_to_text(self, table) -> str:
        """Convert table to text representation"""
        try:
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cells.append(cell_text)
                if any(cells):  # Only add non-empty rows
                    rows.append(" | ".join(cells))
            
            return "\n".join(rows)
            
        except Exception as e:
            logger.warning(f"Failed to convert table to text: {e}")
            return ""
    
    def _create_visual_segments(self, visual_elements: List[VisualElement]) -> List[Segment]:
        """Create segments from visual elements"""
        visual_segments = []
        
        for visual_elem in visual_elements:
            # Determine visual subtype
            if visual_elem.element_type == VisualElementType.IMAGE:
                subtype = VisualSubtype.IMAGE.value
            elif visual_elem.element_type == VisualElementType.CHART:
                subtype = VisualSubtype.CHART.value
            elif visual_elem.element_type == VisualElementType.DIAGRAM:
                subtype = VisualSubtype.DIAGRAM.value
            else:
                subtype = VisualSubtype.IMAGE.value  # Default
            
            # Create placeholder content
            content = f"[{visual_elem.element_type.value.upper()}: Placeholder]"
            
            # Create visual segment
            segment = Segment(
                content=content,
                page_number=visual_elem.page_or_slide,  # paragraph index as page
                segment_type=SegmentType.VISUAL,
                segment_subtype=subtype,
                visual_references=[visual_elem.content_hash],
                metadata={
                    "visual_type": visual_elem.element_type.value,
                    "embed_id": visual_elem.analysis_metadata.get("embed_id", ""),
                    "source": "docx_visual_extraction"
                }
            )
            
            visual_segments.append(segment)
        
        return visual_segments